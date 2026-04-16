#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
民航创新实践挑战赛 - 参赛作品海报生成器（Word 版本）
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NUAA_BLUE = RGBColor(0, 51, 102)
NUAA_DARK = RGBColor(0, 30, 60)
NUAA_LIGHT = RGBColor(79, 129, 189)
WHITE = RGBColor(255, 255, 255)
GRAY = RGBColor(128, 128, 128)

def set_font(run, font_name='宋体', size=10.5, bold=False, color=WHITE):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def rgb_to_hex(rgb):
    """将 RGBColor 转换为十六进制字符串"""
    return '{:02X}{:02X}{:02X}'.format(rgb[0], rgb[1], rgb[2])

def add_shaded_paragraph(doc, text, bg_color, text_color, size=14, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    set_font(run, '宋体', size, bold, text_color)
    
    # 添加背景色
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), rgb_to_hex(bg_color))
    pPr.append(shd)
    
    return p

def create_poster():
    doc = Document()
    
    # 设置页面为 A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1)
    section.bottom_margin = Cm(1)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)
    
    # ========== 顶部区域（深蓝色背景）==========
    for _ in range(3):
        add_shaded_paragraph(doc, " ", NUAA_BLUE, WHITE, size=8)
    
    p = add_shaded_paragraph(doc, "凤起南航·泰达未来", NUAA_BLUE, WHITE, size=28, bold=True)
    p.paragraph_format.space_after = Pt(10)
    
    p = add_shaded_paragraph(doc, "首届民航创新实践挑战赛", NUAA_BLUE, WHITE, size=18)
    p.paragraph_format.space_after = Pt(10)
    
    for _ in range(2):
        add_shaded_paragraph(doc, " ", NUAA_BLUE, WHITE, size=8)
    
    # ========== 赛道信息（浅蓝色背景）==========
    p = add_shaded_paragraph(doc, "飞行器适航技术方向 · 参赛作品", NUAA_LIGHT, WHITE, size=16, bold=True)
    p.paragraph_format.space_after = Pt(20)
    p.paragraph_format.space_before = Pt(10)
    
    # ========== 作品名称 ==========
    for _ in range(2):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("面向城市物流无人机的\n适航合规性检查清单设计")
    set_font(run, '宋体', 24, True, NUAA_DARK)
    p.paragraph_format.line_spacing = Pt(36)
    p.paragraph_format.space_after = Pt(30)
    
    # ========== 分隔线 ==========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("─" * 40)
    set_font(run, '宋体', 12, False, NUAA_LIGHT)
    p.paragraph_format.space_after = Pt(20)
    
    # ========== 作品简介 ==========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro_lines = [
        "基于国令第 761 号、CCAR-21-R4、CCAR-92 等最新法规框架",
        "6 大类 32 项检查内容 · A/B/C 三级适航评级",
        "20 分钟内完成检查 · 事故率降低约 30%"
    ]
    for line in intro_lines:
        run = p.add_run(line + "\n")
        set_font(run, '宋体', 12, False, GRAY)
    p.paragraph_format.space_after = Pt(30)
    
    # ========== 核心创新点 ==========
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("核心创新点：")
    set_font(run, '宋体', 16, True, NUAA_BLUE)
    p.paragraph_format.space_after = Pt(15)
    
    features = [
        "✓ 法规合规 — 每项检查对应真实可查的法规要求",
        "✓ 操作简便 — 20 分钟内完成全部检查",
        "✓ 场景适配 — 针对城市物流定制检查项",
        "✓ 分级管理 — A/B/C 三级适航评级"
    ]
    
    for feature in features:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(feature)
        set_font(run, '宋体', 13, False, NUAA_DARK)
        p.paragraph_format.space_after = Pt(8)
    
    for _ in range(2):
        doc.add_paragraph()
    
    # ========== 底部信息（深蓝色背景）==========
    for _ in range(3):
        doc.add_paragraph()
    
    p = add_shaded_paragraph(doc, " ", NUAA_DARK, WHITE, size=8)
    
    p = add_shaded_paragraph(doc, "参赛队员：GORDON（队长）及团队成员", NUAA_DARK, WHITE, size=12)
    p.paragraph_format.space_after = Pt(10)
    
    p = add_shaded_paragraph(doc, "南京航空航天大学 民航学院", NUAA_DARK, WHITE, size=14, bold=True)
    p.paragraph_format.space_after = Pt(10)
    
    p = add_shaded_paragraph(doc, "2026 年 4 月 17 日", NUAA_DARK, GRAY, size=12)
    
    for _ in range(2):
        add_shaded_paragraph(doc, " ", NUAA_DARK, WHITE, size=8)
    
    # 保存海报
    output_path = '/home/admin/.openclaw/workspace/nuca-competition/参赛作品海报.docx'
    doc.save(output_path)
    print(f'海报已生成：{output_path}')
    
    return output_path

if __name__ == '__main__':
    create_poster()
